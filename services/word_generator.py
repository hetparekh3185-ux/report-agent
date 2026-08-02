import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BOLD_PATTERN = re.compile(r"\*\*(.+?)\*\*")


def _force_font(rpr_owner, font_name: str) -> None:
    """
    Set an explicit font face AND strip any theme-font reference
    (asciiTheme/hAnsiTheme/eastAsiaTheme/csTheme) on the same element.

    This matters because Word's built-in Heading styles ship with a
    w:rFonts element that already has w:asciiTheme="majorHAnsi" (which
    resolves to Calibri Light). python-docx's plain `.font.name = "..."`
    setter only adds w:ascii/w:hAnsi — it does NOT remove the theme
    attributes — and when both an explicit face and a theme reference
    are present on the same rFonts element, Word's renderer prefers the
    theme one. So setting .name alone silently does nothing visible for
    heading styles; the theme attributes have to be deleted too.
    """
    rpr = rpr_owner.element.get_or_add_rPr() if hasattr(rpr_owner, "element") else rpr_owner
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rFonts)
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:csTheme"):
        if rFonts.get(qn(attr)) is not None:
            del rFonts.attrib[qn(attr)]
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rFonts.set(qn(attr), font_name)


def setup_document_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    _force_font(style.element.get_or_add_rPr(), "Times New Roman")

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.name = "Times New Roman"
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        _force_font(heading_style.element.get_or_add_rPr(), "Times New Roman")

        if level == 1:
            heading_style.font.size = Pt(16)
        elif level == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)


def _apply_run_font(run, font_name: str, size_pt: int, bold: bool = None) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    _force_font(run.element.get_or_add_rPr(), font_name)


def _add_text_with_bold_markers(paragraph, text: str, font_name: str, size_pt: int, base_bold: bool = False) -> None:
    """
    Splits `**bold**` markdown into separate runs so it renders as actual
    bold text instead of showing the literal asterisk characters.
    """
    pos = 0
    for match in BOLD_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            _apply_run_font(run, font_name, size_pt, bold=base_bold)
        run = paragraph.add_run(match.group(1))
        _apply_run_font(run, font_name, size_pt, bold=True)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _apply_run_font(run, font_name, size_pt, bold=base_bold)


def parse_and_format_content(doc: Document, content: str) -> None:
    for raw_line in content.split("\n"):
        line = raw_line.strip()
        if not line:
            continue

        # Bullet lines ("- **Foo**: bar") also carry markdown bold, so
        # they go through the same bold-marker splitter as body text.
        if line.startswith("# "):
            heading = doc.add_heading("", level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_text_with_bold_markers(heading, line[2:], "Times New Roman", 16, base_bold=True)
        elif line.startswith("## "):
            heading = doc.add_heading("", level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_text_with_bold_markers(heading, line[3:], "Times New Roman", 14, base_bold=True)
        elif line.startswith("### "):
            heading = doc.add_heading("", level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_text_with_bold_markers(heading, line[4:], "Times New Roman", 12, base_bold=True)
        else:
            para = doc.add_paragraph()
            para.style = doc.styles["Normal"]
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _add_text_with_bold_markers(para, line, "Times New Roman", 11, base_bold=False)


def build_docx(topic: str, body_text: str, output_path: str) -> str:
    doc = Document()
    setup_document_styles(doc)

    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_text_with_bold_markers(title, topic, "Times New Roman", 16, base_bold=True)

    doc.add_paragraph()
    parse_and_format_content(doc, body_text)

    doc.save(output_path)
    return output_path
