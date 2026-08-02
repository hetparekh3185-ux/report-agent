import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from groq import Groq
from dotenv import load_dotenv

load_dotenv()

GROQ_API_KEY = os.getenv('GROQ_API_KEY')

def get_groq_client():
    if not GROQ_API_KEY:
        print('Error: GROQ_API_KEY not found in environment variables.')
        print('Please create a .env file with GROQ_API_KEY=your_api_key')
        sys.exit(1)
    return Groq(api_key=GROQ_API_KEY)

def fetch_content(client, topic, num_pages):
    words_per_page = 350
    target_words = num_pages * words_per_page
    
    prompt = f'''
    Write a comprehensive report on the topic: "{topic}"
    
    Requirements:
    - Target length: approximately {target_words} words ({num_pages} pages)
    - Structure with clear headings and subheadings
    - Include an introduction, main body with multiple sections, and conclusion
    - Use professional, informative tone
    - Include relevant details, examples, and analysis
    - Format with clear section breaks
    '''
    
    try:
        response = client.chat.completions.create(
            model='llama-3.3-70b-versatile',
            messages=[
                {'role': 'system', 'content': 'You are a professional report writer. Create well-structured, informative reports.'},
                {'role': 'user', 'content': prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f'Error fetching content from Groq: {e}')
        sys.exit(1)

def setup_document_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)
    
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
        heading_style.font.bold = True
        
        if level == 1:
            heading_style.font.size = Pt(16)
        elif level == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)

def parse_and_format_content(doc, content):
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            para = doc.add_paragraph(line)
            para.style = doc.styles['Normal']
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)

def create_report(topic, num_pages, output_filename):
    client = get_groq_client()
    
    print(f'Fetching content for topic: "{topic}" ({num_pages} pages)...')
    content = fetch_content(client, topic, num_pages)
    
    doc = Document()
    setup_document_styles(doc)
    
    title = doc.add_heading(topic, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
    
    doc.add_paragraph()
    
    parse_and_format_content(doc, content)
    
    doc.save(output_filename)
    print(f'Report saved as: {output_filename}')

def main():
    print('=' * 50)
    print('Word Report Generator Agent')
    print('=' * 50)
    
    topic = input('Enter the report topic: ').strip()
    if not topic:
        print('Error: Topic cannot be empty.')
        return
    
    try:
        num_pages = int(input('Enter number of pages: ').strip())
        if num_pages <= 0:
            print('Error: Number of pages must be positive.')
            return
    except ValueError:
        print('Error: Please enter a valid number for pages.')
        return
    
    output_filename = input('Enter output filename (default: report.docx): ').strip()
    if not output_filename:
        output_filename = 'report.docx'
    if not output_filename.endswith('.docx'):
        output_filename += '.docx'
    
    create_report(topic, num_pages, output_filename)

if __name__ == '__main__':
    main()
